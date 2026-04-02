# -*- coding: utf-8 -*-
"""Gera BRIEFING_ALTERACOES_CLIENTE.docx a partir da estrutura do briefing."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, fill_hex: str):
    """Cor de fundo da célula (hex sem #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_rich_paragraph(doc, text: str, style=None):
    """Parágrafo com **negrito** simples."""
    p = doc.add_paragraph(style=style)
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True
    return p


def add_table(doc, headers, rows, col_widths_inches=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], "E8E8E8")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val) if val is not None else ""
            for p in tbl.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths_inches:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths_inches):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return tbl


def main():
    out = Path(__file__).resolve().parent / "BRIEFING_ALTERACOES_CLIENTE.docx"
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Capa / título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Briefing de alterações de conteúdo")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = st.add_run("Site — MB Mendanha Sociedade de Advocacia")
    rr.font.size = Pt(14)
    rr.italic = True

    doc.add_paragraph()
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_intro.add_run("Documento para preenchimento pelo cliente").font.size = Pt(10)
    p_intro.runs[0].font.color.rgb = RGBColor(0x68, 0x65, 0x5E)

    doc.add_paragraph("_" * 72)

    doc.add_heading("Sobre este documento", level=2)
    add_table(
        doc,
        ["Tópico", "Descrição"],
        [
            ("Finalidade", "Reunir sugestões de texto e de informações de negócio para o site."),
            (
                "O que não muda por aqui",
                "Estrutura técnica: menus, ordem das seções, fluxo geral das páginas.",
            ),
            (
                "Como preencher",
                "Use as tabelas e campos abaixo. Se algo deve permanecer igual, escreva «manter como está» ou deixe em branco.",
            ),
        ],
        [1.4, 4.8],
    )

    doc.add_heading("Índice", level=2)
    idx = doc.add_paragraph()
    items = [
        "1. Dados institucionais",
        "2. SEO e redes (textos)",
        "3. Página inicial",
        "4. Fluxo WhatsApp",
        "5. Páginas internas",
        "6. Rodapé e informações legais",
        "7. Tom de voz e público",
        "8. Fora do escopo deste briefing",
        "9. Outras observações",
        "10. Aprovação",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Number")

    doc.add_page_break()

    # --- Seção 1 ---
    doc.add_heading("1. Dados institucionais", level=1)
    doc.add_heading("1.1 Identificação e contatos", level=2)
    add_table(
        doc,
        ["Campo", "Conteúdo atual (referência)", "Sua sugestão / novo texto"],
        [
            ("Nome do escritório", "MB Mendanha Sociedade de Advocacia", ""),
            ("Advogado(a) responsável", "Marcos Aurélio Mendanha", ""),
            ("OAB", "OAB/SP 368.689", ""),
            ("E-mail", "advocacia.mendanha@gmail.com", ""),
            ("WhatsApp", "(11) 96503-0211", ""),
            ("Instagram", "@mbmendanha.advocacia", ""),
        ],
        [1.5, 2.5, 2.8],
    )

    doc.add_heading("1.2 Endereço (rodapé, contato, mapas)", level=2)
    add_table(
        doc,
        ["", "Conteúdo atual", "Sua sugestão"],
        [
            ("Logradouro e complemento", "Av. Imirim, 4301 - SALA 03", ""),
            ("Bairro", "Vila Nova Cachoeirinha", ""),
            ("Cidade / UF", "São Paulo – SP", ""),
            ("CEP", "02465-500", ""),
        ],
        [1.6, 2.4, 2.8],
    )

    # --- Seção 2 ---
    doc.add_heading("2. SEO e redes (textos)", level=1)
    add_rich_paragraph(
        doc,
        "Estes textos aparecem em **resultados de busca** e ao **compartilhar o link**. "
        "A estrutura técnica da página permanece a mesma; aqui você só indica **revisão de texto** ou **novas palavras-chave**, se desejar.",
    )
    add_table(
        doc,
        ["Item", "Nota", "Sua sugestão"],
        [
            ("Título (aba do navegador)", "Hoje: Civil, Família, Trabalho e Zona Norte SP", ""),
            ("Descrição curta (snippet)", "Cerca de 1 a 2 frases", ""),
            ("Palavras-chave", "Separadas por vírgula", ""),
        ],
        [1.6, 2.4, 2.6],
    )

    # --- Seção 3 ---
    doc.add_heading("3. Página inicial", level=1)
    doc.add_paragraph(
        "Para cada bloco, indique «manter» ou cole/descreva o novo texto."
    )
    add_table(
        doc,
        ["#", "Bloco no site", "O que revisar", "Suas observações / novo texto"],
        [
            ("1", "Faixa acima do título (hero)", "Nome do escritório + região", ""),
            ("2", "Título principal (H1)", "Frase de impacto principal", ""),
            ("3", "Parágrafo do hero", "Apoio ao título", ""),
            ("4", "Seção «Sobre o escritório»", "Texto institucional", ""),
            ("5", "Citação em destaque", "Frase atribuída ao advogado", ""),
            ("6", "Seção «Áreas de atuação»", "Título + parágrafo introdutório", ""),
            (
                "7",
                "Três cards (Civil / Família / Trabalhista)",
                "Título, texto e textos dos botões WhatsApp",
                "",
            ),
            ("8", "«Como funciona»", "Os 4 passos", ""),
            ("9", "«Atendimento local»", "Texto + lista de bairros", ""),
            ("10", "FAQ resumido (na home)", "Perguntas e respostas", ""),
            ("11", "Chamada final (antes dos botões)", "CTA", ""),
        ],
        [0.35, 1.5, 1.5, 2.3],
    )

    # --- Seção 4 ---
    doc.add_heading("4. Fluxo WhatsApp", level=1)
    add_rich_paragraph(
        doc,
        "A **ordem das etapas** é fixa no site: **(1) área → (2) tema → (3) descrição**. "
        "Abaixo você pode sugerir nomes dos botões, lista de temas e textos auxiliares.",
    )
    doc.add_heading("4.1 Nomes dos botões de área", level=2)
    add_rich_paragraph(doc, "Valores atuais: **Civil** · **Família** · **Trabalhista**")
    doc.add_paragraph("Deseja alterar algum nome? (sim / não)")
    doc.add_paragraph("Se sim, descreva:")
    doc.add_paragraph("_" * 60)

    doc.add_heading("4.2 Subtemas (menu suspenso)", level=2)
    add_table(
        doc,
        ["Área", "Temas atuais", "Incluir / remover / substituir por"],
        [
            ("Civil", "Contratos, Indenização, Cobrança, Outro", ""),
            ("Família", "Divórcio, Pensão, Guarda, Outro", ""),
            ("Trabalhista", "Demissão, Horas extras, FGTS, Assédio, Outro", ""),
        ],
        [1.2, 2.3, 2.9],
    )

    doc.add_heading("4.3 Mensagem automática e campo de descrição", level=2)
    add_table(
        doc,
        ["Campo", "Situação hoje", "Sua sugestão"],
        [
            (
                "Abertura da mensagem no WhatsApp",
                "Inclui o nome do escritório, área, tema e descrição informada pelo visitante",
                "Frase de abertura diferente? (descreva)",
            ),
            (
                "Texto de exemplo na caixa «descreva o caso»",
                "Placeholder no site",
                "Novo texto de exemplo:",
            ),
        ],
        [1.8, 2.2, 2.6],
    )

    # --- Seção 5 ---
    doc.add_heading("5. Páginas internas", level=1)
    doc.add_paragraph(
        "O site reúne várias «páginas» no mesmo arquivo; abaixo, o que costuma ser editado em cada uma."
    )
    add_table(
        doc,
        ["Página", "Conteúdos típicos", "Suas observações"],
        [
            ("Sobre", "Título, história, bio, valores", ""),
            ("Áreas (detalhada)", "Subtítulos e textos por área (Civil, Família, Trabalho)", ""),
            (
                "FAQ",
                "Perguntas e respostas (alterações jurídicas: revisar com advogado)",
                "",
            ),
            ("Contato", "Texto de apoio, horário, bloco de atendimento online", ""),
        ],
        [1.1, 2.8, 2.5],
    )

    # --- Seção 6 ---
    doc.add_heading("6. Rodapé e informações legais", level=1)
    add_table(
        doc,
        ["Item", "Referência atual", "Alteração desejada"],
        [
            ("Texto resumo do escritório", "Direito Civil, Família e Trabalho…", ""),
            ("Links da coluna «Áreas»", "Civil, Família, Trabalho, WhatsApp", ""),
            ("Aviso legal (site informativo)", "Texto padrão de compliance", "manter / ajustar (descreva)"),
            ("Ano do copyright", "2026", ""),
        ],
        [1.6, 2.4, 2.6],
    )

    # --- Seção 7 ---
    doc.add_heading("7. Tom de voz e público", level=1)
    doc.add_paragraph("Marque o que mais se aproxima do tom desejado (pode marcar mais de um):")
    for line in [
        "Mais formal",
        "Mais acessível e direto",
        "Ênfase em acolhimento / atendimento humanizado",
        "Ênfase em experiência e resultados",
    ]:
        doc.add_paragraph(f"☐ {line}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Público principal ").bold = True
    p.add_run("(ex.: trabalhadores da Zona Norte, famílias, empresas):")
    doc.add_paragraph("_" * 60)

    p2 = doc.add_paragraph()
    p2.add_run("Conteúdos ou promessas que não devem aparecer no site ").bold = True
    p2.add_run("(ex.: valores fixos, garantia de resultado):")
    doc.add_paragraph("_" * 60)

    # --- Seção 8 ---
    doc.add_heading("8. Fora do escopo deste briefing", level=1)
    add_rich_paragraph(
        doc,
        "Os itens abaixo costumam exigir **trabalho de desenvolvimento, design ou integração**, "
        "não apenas troca de texto neste formulário:",
    )
    for bullet in [
        "Alteração de layout, cores, fontes ou animações",
        "Criação de páginas ou seções novas",
        "Formulários que enviem e-mail automaticamente (servidor / API)",
        "Fotos, vídeos ou manual de identidade visual",
        "Domínio, hospedagem, Google Meu Negócio",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")
    doc.add_paragraph(
        "Se mesmo assim precisar de algo dessa lista, descreva na seção 9 (Outras observações)."
    )

    # --- Seção 9 ---
    doc.add_heading("9. Outras observações", level=1)
    doc.add_paragraph("Prioridades, prazos, dúvidas ou sites de referência.")
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()
    doc.add_paragraph("_" * 60)

    # --- Seção 10 ---
    doc.add_heading("10. Aprovação", level=1)
    add_table(
        doc,
        ["Campo", "Preenchimento"],
        [
            ("Nome", ""),
            ("Data", ""),
            ("E-mail ou telefone", ""),
        ],
        [1.8, 4.4],
    )

    doc.add_heading("Nota final", level=2)
    add_rich_paragraph(
        doc,
        "Documento para orientar o preenchimento pelo cliente. Conteúdos com implicação **jurídica** "
        "ou de **publicidade profissional** devem ser validados pelo **advogado responsável**.",
    )

    # Rodapé com numeração de páginas (opcional - precisa section)
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    if not footer.paragraphs:
        fp = footer.add_paragraph()
    else:
        fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Briefing de alterações — MB Mendanha Sociedade de Advocacia")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(out))
    print(f"Gerado: {out}")


if __name__ == "__main__":
    main()
